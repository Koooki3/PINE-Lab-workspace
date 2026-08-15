from __future__ import annotations
import argparse, json, re
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

FILES=['00-index.md','01-foundations.md','02-reinforcement-learning.md','03-robot-learning.md','04-paper-guides.md','05-knowledge-graph.md','06-source-ledger.md','07-learning-path.md']
COLORS={1:'#DCEBFA',2:'#DFF3E4',3:'#FFF0C9',4:'#F8D8D8',5:'#E7DDF6'}

def font(size):
    candidates=[Path(r'C:\Windows\Fonts\msyh.ttc'),Path(r'C:\Windows\Fonts\simhei.ttf')]
    for p in candidates:
        if p.exists(): return ImageFont.truetype(str(p),size)
    return ImageFont.load_default()

def graph_png(graph, out):
    W,H=1800,950; img=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(img); title=font(34); body=font(23); small=font(18)
    d.text((55,25),'PINE References 知识图谱：从基础到真实机器人闭环',font=title,fill='#17324D')
    levels={i:[] for i in range(1,6)}
    for n in graph['nodes']: levels[n['level']].append(n)
    pos={}
    for lv,nodes in levels.items():
        x=45+(lv-1)*350; gap=790/max(1,len(nodes))
        for j,n in enumerate(nodes): pos[n['id']]=(x,105+j*gap,290,72)
    for e in graph['edges']:
        x1,y1,w1,h1=pos[e['source']]; x2,y2,w2,h2=pos[e['target']]
        d.line((x1+w1,y1+h1/2,x2,y2+h2/2),fill='#9AA8B5',width=3)
        d.polygon([(x2,y2+h2/2),(x2-12,y2+h2/2-6),(x2-12,y2+h2/2+6)],fill='#9AA8B5')
    for lv,nodes in levels.items():
        for n in nodes:
            x,y,w,h=pos[n['id']]; d.rounded_rectangle((x,y,x+w,y+h),14,fill=COLORS[lv],outline='#52677A',width=2)
            label=n['label']; bbox=d.textbbox((0,0),label,font=body); d.text((x+(w-(bbox[2]-bbox[0]))/2,y+20),label,font=body,fill='#172B3A')
        d.text((45+(lv-1)*350,885),f'L{lv}',font=small,fill='#52677A')
    img.save(out,dpi=(180,180))

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)

def set_run(run,size=10.5,bold=False,color=None):
    run.font.name='Aptos'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); run.font.size=Pt(size); run.bold=bold
    if color: run.font.color.rgb=RGBColor(*color)

def add_markdown(doc,text,skip_title=False):
    lines=text.splitlines(); i=0; first_heading=True; in_code=False
    while i<len(lines):
        line=lines[i].rstrip()
        if line.startswith('```'): in_code=not in_code; i+=1; continue
        if in_code: i+=1; continue
        if line.startswith('|') and i+1<len(lines) and re.match(r'^\|[\s|:-]+\|$',lines[i+1]):
            rows=[]; i+=2
            header=[c.strip() for c in line.strip('|').split('|')]
            while i<len(lines) and lines[i].startswith('|'):
                rows.append([c.strip() for c in lines[i].strip('|').split('|')]); i+=1
            t=doc.add_table(rows=1,cols=len(header)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
            for j,v in enumerate(header): shade(t.rows[0].cells[j],'D9EAF7'); r=t.rows[0].cells[j].paragraphs[0].add_run(v); set_run(r,9,bold=True)
            for row in rows:
                cells=t.add_row().cells
                for j,v in enumerate(row): cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; r=cells[j].paragraphs[0].add_run(v); set_run(r,8.5)
            continue
        m=re.match(r'^(#{1,3})\s+(.*)',line)
        if m:
            if skip_title and first_heading: first_heading=False; i+=1; continue
            p=doc.add_heading(m.group(2),level=min(len(m.group(1)),3)); first_heading=False
            i+=1; continue
        if line.startswith('- '):
            p=doc.add_paragraph(style='List Bullet'); r=p.add_run(line[2:]); set_run(r); i+=1; continue
        if re.match(r'^\d+\.\s',line):
            p=doc.add_paragraph(style='List Number'); r=p.add_run(re.sub(r'^\d+\.\s','',line)); set_run(r); i+=1; continue
        if line.strip():
            p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0.74); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
            chunks=re.split(r'(\*\*.*?\*\*|`.*?`)',line)
            for c in chunks:
                if not c: continue
                bold=c.startswith('**'); code=c.startswith('`') and c.endswith('`'); val=c[2:-2] if bold else c[1:-1] if code else c
                r=p.add_run(val); set_run(r,10.2,bold=bold,color=(80,60,120) if code else None)
        i+=1

def page_field(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=paragraph.add_run('— '); set_run(run,9,color=(100,110,120))
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld); run=paragraph.add_run(' —'); set_run(run,9,color=(100,110,120))

def build(root):
    kg=root/'references'/'Knowledge graph'; graph=json.loads((kg/'graph.json').read_text(encoding='utf-8')); graph_path=kg/'knowledge-graph.png'; graph_png(graph,graph_path)
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(2); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.2)
    styles=doc.styles
    styles['Normal'].font.name='Aptos'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
    for name,size,color in [('Title',28,(23,50,77)),('Heading 1',19,(23,50,77)),('Heading 2',14,(35,92,126)),('Heading 3',11.5,(68,110,125))]:
        s=styles[name]; s.font.name='Aptos Display'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); s.font.size=Pt(size); s.font.color.rgb=RGBColor(*color)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110)
    r=p.add_run('PINE'); set_run(r,34,bold=True,color=(23,50,77)); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('机器人学习 References 知识库与知识图谱'); set_run(r,21,bold=True,color=(35,92,126))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('从强化学习基础到通用真实机器人控制'); set_run(r,13,color=(90,105,115))
    doc.add_paragraph(); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('仓库内论文驱动 · 前置知识补全 · 易到难学习路线'); set_run(r,10,color=(90,105,115))
    page_field(sec.footer.paragraphs[0]); doc.add_page_break()
    doc.add_heading('如何使用本手册',0); add_markdown(doc,(kg/'00-index.md').read_text(encoding='utf-8'),skip_title=True)
    doc.add_page_break(); doc.add_heading('全局知识图谱',0); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(graph_path),width=Inches(6.7)); doc.add_paragraph('箭头表示“前者为后者提供前置、能力或方法来源”；L1-L5 表示由基础到综合研究假设。').alignment=WD_ALIGN_PARAGRAPH.CENTER
    for f in FILES[1:]:
        doc.add_page_break(); add_markdown(doc,(kg/f).read_text(encoding='utf-8'))
    props=doc.core_properties; props.title='PINE机器人学习References知识库与知识图谱'; props.subject='SAC, π0, Q-chunking, HIL-SERL'; props.author='PINE Lab'
    out=kg/'PINE.docx'; doc.save(out); graph_path.unlink(); print(out)

if __name__=='__main__':
    ap=argparse.ArgumentParser(); ap.add_argument('--root',required=True); a=ap.parse_args(); build(Path(a.root).resolve())
